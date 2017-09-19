from docx import Document
from docx.shared import Inches
import scrapy
import re
import csv
from scrapy.http.request import Request
from scrapy.selector import Selector

# Create object for word document
document = Document()

class WenxuecitySpider(scrapy.Spider):
    name = "WenxuecityWord"

    def start_requests(self):
        # Open csv file for reading purposes
        with open('./csv/wenxuecity.csv') as ffile:
            reader = csv.reader(ffile, delimiter=',')
            for row in reader:
                request = Request(row[3], callback=self.parse_link)
                request.meta['type'] = row[4]
                request.meta['id'] = row[0]
                yield request

    def parse_link(self, response):

        # Remove random <br> tags for better organization
        response = response.replace(body=response.body.replace(b'<br>', b'\n'))
        response = response.replace(body=response.body.replace(b'\n', b''))
        sel = Selector(response)

        # Content
        content = sel.xpath('//*[@id="articleContent"]/p/text()')
        content = content.extract()

        if (len(content) == 0):
            # Xpath selector didn't match; try another Xpath selector
            content = sel.xpath('//*[@id="articleContent"]/text()')
            content = content.extract()

        full_content = ""
        for paragraph in content:
            full_content += paragraph

        full_content = full_content.encode('unicode-escape').decode('unicode-escape')
        document.add_paragraph(response.meta['id'] + '\n')
        document.add_paragraph(full_content)
        document.save('./word/wenxuecity.docx')
