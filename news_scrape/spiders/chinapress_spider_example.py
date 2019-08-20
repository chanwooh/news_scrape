from docx import Document
from docx.shared import Inches
import scrapy
import re
import csv
from scrapy.http.request import Request
from scrapy.selector import Selector

# Create object for word document
document = Document()

class ChinaPressSpider(scrapy.Spider):
    name = "SingtaoExample"

    def start_requests(self):
        request = Request('https://www.singtaousa.com/la/453-%E5%8D%97%E5%8A%A0%E6%96%B0%E8%81%9E/352233-%E4%B8%AD%E7%BE%8E%E9%9B%BB%E8%A6%96%E5%8A%87%E9%A0%90%E7%AE%97%E5%B7%AE%E8%B7%9D%E6%B8%9B+%E5%B0%A4%E5%B0%8F%E5%89%9B%EF%BC%9A%E5%90%88%E6%8B%8D%E6%A9%9F%E7%8E%87%E5%A2%9E%E5%8A%A0/?fromG=1', callback=self.parse_link)
        request.meta['type'] = 'ST'
        request.meta['id'] = 'lit'
        yield request
                

    def parse_link(self, response):

        # Remove random <br> tags for better organization
        response = response.replace(body=response.body.replace(b'<br>', b'\n'))
        response = response.replace(body=response.body.replace(b'\n', b''))

        sel = Selector(response)

        # Content
        content = sel.xpath('//*[@id="mobTouchSwipeWrap"]/div[1]/div/div[1]/div[2]/div[4]/p/text()')
        content = content.extract()

        full_content = ""
        for paragraph in content:
            full_content += paragraph

        full_content = full_content.encode('unicode-escape').decode('unicode-escape')
        document.add_paragraph(response.meta['id'] + '\n')
        document.add_paragraph(full_content)
        document.save('./word/singtao.docx')
