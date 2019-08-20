# -*- coding: utf-8 -*-
import scrapy
import re
import json
import docx
import csv
from scrapy.http.request import Request
from scrapy.selector import Selector
from docx import Document
from docx.shared import Inches

# Create object for word document
document = Document()

class AggregateSpider(scrapy.Spider):
    name = "AggregateWord"

    def start_requests(self):

        # Read from CSV file and make requests
        with open('./csv/aggregate.csv') as ffile:
            reader = csv.reader(ffile, delimiter=',')
            for row in reader:
                request = Request(url=row[4], callback=self.parse_link)
                request.meta['id'] = row[0]
                yield request

    def parse_link(self, response):

        # Remove random <br> tags for better organization
        response = response.replace(body=response.body.replace(b'<br>', b'\n'))
        response = response.replace(body=response.body.replace(b'\n', b''))

        sel = Selector(response)

        # Content
        content = sel.xpath('//*[@id="js_content"]//p//text()')
        content = content.extract()

        full_content = ""
        for paragraph in content:
            full_content += paragraph

        full_content = full_content.encode('unicode-escape').decode('unicode-escape')
        p = document.add_paragraph('')
        p.add_run(response.meta['id'] + '\n').bold = True

        if (full_content != ""):
            document.add_paragraph(full_content)
        else:
            document.add_paragraph("Content is unavailable. It has been deleted, moved, or requires a QR scan.")

        document.save('./word/aggregate.docx')
